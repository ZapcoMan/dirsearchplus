#!/usr/bin/env python3
# -*- encoding: utf-8 -*-

import os,sqlite3,time,json
from docx import Document   #用来建立一个word对象
from copy import deepcopy
from docx.shared import Pt  #用来设置字体的大小
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import RGBColor
from urllib.parse import urlparse
from lib.common.utils import Utils
from lib.Database import DatabaseType
from lib.common.CreatLog import creatLog
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from lib.common.cmdline import CommandLines


class Creat_vuln_detail():
    """
    漏洞详情文档生成类

    用于根据数据库中的漏洞信息，在Word文档中自动生成详细的漏洞描述、接口地址、JS路径及响应内容等信息。

    Attributes:
        projectTag (str): 项目标识符，用于定位对应的数据库文件
        creat_num (int): 编号计数器（正向）
        creat_num1 (int): 编号计数器（反向）
        log: 日志记录器实例
    """

    def __init__(self,projectTag):
        """
        初始化Creat_vuln_detail类

        Args:
            projectTag (str): 项目的唯一标识标签
        """
        self.projectTag = projectTag
        self.creat_num = 1
        self.creat_num1 = 1
        self.log = creatLog().get_logger()

    def copy_table_after(self,table, paragraph):
        """
        在指定段落后复制表格

        Args:
            table: 要复制的表格对象
            paragraph: 目标段落对象
        """
        tbl, p = table._tbl, paragraph._p
        new_tbl = deepcopy(tbl)
        p.addnext(new_tbl)

    def insert_paragraph_after(self,paragraph, text=None, style=None):
        """
        在指定段落后插入新段落

        Args:
            paragraph: 参考段落对象
            text (str, optional): 新段落的文本内容
            style (str, optional): 新段落的样式

        Returns:
            Paragraph: 新创建的段落对象
        """
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        if text:
            new_para.add_run(text)
        if style is not None:
            new_para.style = style
        return new_para

    def move_table_after(self, table, paragraph):
        """
        将表格移动到指定段落之后

        Args:
            table: 要移动的表格对象
            paragraph: 目标段落对象
        """
        tbl, p = table._tbl, paragraph._p
        p.addnext(tbl)

    def tabBgColor(self,table, cols, colorStr):
        """
        设置表格首行背景色

        Args:
            table: 表格对象
            cols (int): 列数
            colorStr (str): 颜色值字符串（如'F5F5F5'）
        """
        shading_list = locals()
        for i in range(cols):
            shading_list['shading_elm_' + str(i)] = parse_xml(
                r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStr))
            table.rows[0].cells[i]._tc.get_or_add_tcPr().append(shading_list['shading_elm_' + str(i)])

    def locat_detail(self,document):
        """
        定位并处理文档中的占位符{vuln_deta}

        Args:
            document: Word文档对象

        Returns:
            Paragraph: 插入的新段落对象
        """
        for para in document.paragraphs:
            for i in range(len(para.runs)):
                if "{vuln_deta}" in para.runs[i].text:
                    para.runs[i].text = para.runs[i].text.replace('{vuln_deta}', '')
                    para1 = para.insert_paragraph_before("")
        return para1

    def creat_table(self,document,vuln_info,para2):
        """
        创建带背景色的表格并插入到指定段落之后

        Args:
            document: Word文档对象
            vuln_info (str): 表格内容
            para2: 参考段落对象
        """
        colorStr = 'F5F5F5'
        row = col = 1
        table = document.add_table(row, col)
        table.rows[0].cells[0].text = "%s" % (vuln_info)
        Creat_vuln_detail(self.projectTag).tabBgColor(table, col, colorStr)
        Creat_vuln_detail(self.projectTag).move_table_after(table, para2)

    def creat_detail(self,document):
        """
        根据数据库漏洞信息生成完整的漏洞详情文档

        从数据库读取各种类型的漏洞信息（未授权访问、信息泄露、CORS配置错误等），
        并按照不同漏洞类型格式化输出到Word文档中。

        Args:
            document: Word文档对象
        """
        # 定位文档中的漏洞详情占位符位置
        para1 = Creat_vuln_detail(self.projectTag).locat_detail(document)

        # 连接项目数据库获取漏洞数据
        projectDBPath = DatabaseType(self.projectTag).getPathfromDB() + self.projectTag + ".db"
        connect = sqlite3.connect(os.sep.join(projectDBPath.split('/')))
        cursor = connect.cursor()
        connect.isolation_level = None
        sql = "select * from vuln"
        cursor.execute(sql)
        vuln_infos = cursor.fetchall()
        k = len(vuln_infos)

        try:
            # 处理各类漏洞信息
            for vuln_info in vuln_infos:
                # 处理确定性未授权访问漏洞
                if vuln_info[3] == "unAuth" and vuln_info[4] == 1:
                    sql = "select * from api_tree where id='%s'" % (vuln_info[1])
                    cursor.execute(sql)
                    api_infos = cursor.fetchall()
                    for api_info in api_infos:
                        para2 = para1.insert_paragraph_before("")
                        UserLogin = api_info[2]
                        api_path = api_info[1]
                        run = para2.add_run("2."+ str(self.creat_num) + " " + str(UserLogin) + Utils().getMyWord("{r_vuln_unauth}") + "\n")
                        run.font.name = "Arial"
                        run.font.size = Pt(16)
                        run.font.bold = True
                        sql = "select path from js_file where id='%s'" % (vuln_info[2])
                        cursor.execute(sql)
                        js_paths = cursor.fetchall()
                        for js_path in js_paths:
                            run2 = para2.add_run(Utils().getMyWord("{r_api_addr}"))
                            run2.font.name = "Arial"
                            run2.font.size = Pt(10)
                            run2.font.bold = True
                            run3 = para2.add_run(api_path)
                            run3.font.name = "Arial"
                            run3.font.size = Pt(10)
                            run4 = para2.add_run("\n" + Utils().getMyWord("{r_api_js}"))
                            run4.font.name = "Arial"
                            run4.font.size = Pt(10)
                            run4.font.bold = True
                            run5 = para2.add_run(js_path[0])
                            run5.font.name = "Arial"
                            run5.font.size = Pt(10)
                            run5 = para2.add_run("\n" + Utils().getMyWord("{r_api_res}"))
                            run5.font.name = "Arial"
                            run5.font.size = Pt(10)
                            run5.font.bold = True
                            self.creat_num = self.creat_num + 1
                            vuln_info_js_unicode = json.dumps(json.loads(vuln_info[6]), sort_keys=True, indent=4,ensure_ascii=False)
                            Creat_vuln_detail(self.projectTag).creat_table(document, vuln_info_js_unicode, para2)

                # 处理可能的未授权访问漏洞
                elif vuln_info[3] == "unAuth" and vuln_info[4] == 2:
                    sql = "select * from api_tree where id='%s'" % (vuln_info[1])
                    cursor.execute(sql)
                    api_infos = cursor.fetchall()
                    for api_info in api_infos:
                        para2 = para1.insert_paragraph_before("")
                        UserLogin = api_info[2]
                        api_path = api_info[1]
                        run = para2.add_run("2." + str(self.creat_num) + " " + str(UserLogin) + Utils().getMyWord("{r_vuln_unauth_maybe}") + "\n")
                        run.font.name = "Arial"
                        run.font.size = Pt(16)
                        run.font.bold = True
                        sql = "select path from js_file where id='%s'" % (vuln_info[2])
                        cursor.execute(sql)
                        js_paths = cursor.fetchall()
                        for js_path in js_paths:
                            run2 = para2.add_run(Utils().getMyWord("{r_api_addr}"))
                            run2.font.name = "Arial"
                            run2.font.size = Pt(10)
                            run2.font.bold = True
                            run3 = para2.add_run(api_path)
                            run3.font.name = "Arial"
                            run3.font.size = Pt(10)
                            run4 = para2.add_run("\n" + Utils().getMyWord("{r_api_js}"))
                            run4.font.name = "Arial"
                            run4.font.size = Pt(10)
                            run4.font.bold = True
                            run5 = para2.add_run(js_path[0])
                            run5.font.name = "Arial"
                            run5.font.size = Pt(10)
                            run5 = para2.add_run("\n" + Utils().getMyWord("{r_api_res}"))
                            run5.font.name = "Arial"
                            run5.font.size = Pt(10)
                            run5.font.bold = True
                            self.creat_num = self.creat_num + 1
                            vuln_info_js_unicode = json.dumps(json.loads(vuln_info[6]), sort_keys=True, indent=4, ensure_ascii=False)
                            Creat_vuln_detail(self.projectTag).creat_table(document, vuln_info_js_unicode, para2)

                # 处理信息泄露漏洞
                elif vuln_info[3] == "INFO":
                    sql = "select * from js_file where id='%s'" % (vuln_info[2])
                    cursor.execute(sql)
                    js_infos = cursor.fetchall()
                    for js_info in js_infos:
                        js_name = js_info[1]
                        js_path = js_info[2]
                        para2 = para1.insert_paragraph_before("")
                        run = para2.add_run("2." + str(self.creat_num) + " " + str(js_name) + Utils().getMyWord("{r_vuln_info}") + "\n")
                        run.font.name = "Arial"
                        run.font.size = Pt(16)
                        run.font.bold = True
                        run2 = para2.add_run(Utils().getMyWord("{r_js_path}"))
                        run2.font.name = "Arial"
                        run2.font.size = Pt(10)
                        run2.font.bold = True
                        run3 = para2.add_run(js_path)
                        run3.font.name = "Arial"
                        run3.font.size = Pt(10)
                        run4 = para2.add_run("\n" + Utils().getMyWord("{r_js_des}"))
                        run4.font.name = "Arial"
                        run4.font.size = Pt(10)
                        run4.font.bold = True
                        run5 = para2.add_run(vuln_info[8])
                        run5.font.name = "Arial"
                        run5.font.size = Pt(10)
                        run6 = para2.add_run("\n" + Utils().getMyWord("{r_js_detial}"))
                        run6.font.name = "Arial"
                        run6.font.size = Pt(10)
                        run6.font.bold = True
                        self.creat_num = self.creat_num + 1
                        Creat_vuln_detail(self.projectTag).creat_table(document, vuln_info[7], para2)

                # 处理CORS跨域配置错误
                elif vuln_info[3] == "CORS":
                    sql = "select vaule from info where name='%s'" % ("host")
                    cursor.execute(sql)
                    infos = cursor.fetchall()
                    for info in infos:
                        api_path = info[0]
                    para2 = Creat_vuln_detail(self.projectTag).insert_paragraph_after(para1)
                    para3 = para2.insert_paragraph_before("")
                    run5 = para3.add_run("2." + str(self.creat_num) + " " + str(api_path) + Utils().getMyWord("{r_vuln_CORS}") + "\n")
                    run5.font.name = "Arial"
                    run5.font.size = Pt(16)
                    run5.font.bold = True
                    run6 = para3.add_run("网址:")
                    run6.font.name = "Arial"
                    run6.font.size = Pt(10)
                    run6.font.bold = True
                    run7 = para3.add_run(api_path)
                    run7.font.name = "Arial"
                    run7.font.size = Pt(10)
                    run8 = para3.add_run("\n" + "{response_head}")
                    run8.font.name = "Arial"
                    run8.font.size = Pt(10)
                    run8.font.bold = True
                    Creat_vuln_detail(self.projectTag).creat_table(document, vuln_info[7], para2)
                    run9 = para2.add_run("\n" + "{request_head}")
                    run9.font.name = "Arial"
                    run9.font.size = Pt(10)
                    run9.font.bold = True
                    Creat_vuln_detail(self.projectTag).creat_table(document, vuln_info[5], para3)

            # 处理第二轮漏洞信息（密码相关、权限控制等）
            for vuln_info in vuln_infos:
                # 处理弱密码漏洞
                if vuln_info[3] == "passWord":
                    sql = "select * from api_tree where id='%s'" % (vuln_info[1])
                    cursor.execute(sql)
                    api_infos = cursor.fetchall()
                    for api_info in api_infos:
                        para2 = Creat_vuln_detail(self.projectTag).insert_paragraph_after(para1)
                        para3 = para2.insert_paragraph_before("")
                        UserLogin = api_info[2]
                        api_path = api_info[1]
                        run = para3.add_run("2." + str(k - self.creat_num1 + 1) + " " + str(UserLogin) + Utils().getMyWord(
                            "{r_vuln_passWord}") + "\n")
                        run.font.name = "Arial"
                        run.font.size = Pt(16)
                        run.font.bold = True
                        sql = "select path from js_file where id='%s'" % (vuln_info[2])
                        cursor.execute(sql)
                        js_paths = cursor.fetchall()
                        for js_path in js_paths:
                            run2 = para3.add_run(Utils().getMyWord("{r_api_addr}"))
                            run2.font.name = "Arial"
                            run2.font.size = Pt(10)
                            run2.font.bold = True
                            run3 = para3.add_run(api_path)
                            run3.font.name = "Arial"
                            run3.font.size = Pt(10)
                            run4 = para3.add_run("\n" + Utils().getMyWord("{r_api_js}"))
                            run4.font.name = "Arial"
                            run4.font.size = Pt(10)
                            run4.font.bold = True
                            run5 = para3.add_run(js_path[0])
                            run5.font.name = "Arial"
                            run5.font.size = Pt(10)
                            run5 = para3.add_run("\n" + Utils().getMyWord("{r_api_res}"))
                            run5.font.name = "Arial"
                            run5.font.size = Pt(10)
                            run5.font.bold = True
                            self.creat_num1 = self.creat_num1 + 1
                            vuln_info_js_unicode = json.dumps(json.loads(vuln_info[6]), sort_keys=True, indent=4,
                                                              ensure_ascii=False)
                            Creat_vuln_detail(self.projectTag).creat_table(document, vuln_info_js_unicode, para2)
                            run6 = para2.add_run(Utils().getMyWord("{request_info}"))
                            run6.font.name = "Arial"
                            run6.font.size = Pt(10)
                            run6.font.bold = True
                            vuln_info_js_unicode = json.dumps(json.loads(vuln_info[5]), sort_keys=True, indent=4,
                                                              ensure_ascii=False)
                            Creat_vuln_detail(self.projectTag).creat_table(document, vuln_info_js_unicode, para3)

                # 处理越权访问漏洞
                elif vuln_info[3] == "BAC":
                    sql = "select * from api_tree where id='%s'" % (vuln_info[1])
                    cursor.execute(sql)
                    api_infos = cursor.fetchall()
                    for api_info in api_infos:
                        para2 = Creat_vuln_detail(self.projectTag).insert_paragraph_after(para1)
                        para3 = para2.insert_paragraph_before("")
                        UserLogin = api_info[2]
                        api_path = api_info[1]
                        run = para3.add_run("2." + str(k - self.creat_num1 + 1) + " " + str(UserLogin) + Utils().getMyWord("{r_vuln_bac}") + "\n")
                        run.font.name = "Arial"
                        run.font.size = Pt(16)
                        run.font.bold = True
                        sql = "select path from js_file where id='%s'" % (vuln_info[2])
                        cursor.execute(sql)
                        js_paths = cursor.fetchall()
                        for js_path in js_paths:
                            run2 = para3.add_run(Utils().getMyWord("{r_api_addr}"))
                            run2.font.name = "Arial"
                            run2.font.size = Pt(10)
                            run2.font.bold = True
                            run3 = para3.add_run(api_path)
                            run3.font.name = "Arial"
                            run3.font.size = Pt(10)
                            run4 = para3.add_run("\n" + Utils().getMyWord("{r_api_js}"))
                            run4.font.name = "Arial"
                            run4.font.size = Pt(10)
                            run4.font.bold = True
                            run5 = para3.add_run(js_path[0])
                            run5.font.name = "Arial"
                            run5.font.size = Pt(10)
                            run5 = para3.add_run("\n" + Utils().getMyWord("{request_info}"))
                            run5.font.name = "Arial"
                            run5.font.size = Pt(10)
                            run5.font.bold = True
                            self.creat_num1 = self.creat_num1 + 1
                            info1 = "请求内容1: " + vuln_info[5].split("§§§")[0] + "\n\n" + "请求内容2: " + vuln_info[5].split("§§§")[1]
                            info2 = "响应内容1: " + vuln_info[6].split("§§§")[0] + "\n\n" + "响应内容2: " + vuln_info[6].split("§§§")[1]
                            Creat_vuln_detail(self.projectTag).creat_table(document, info2 ,para2)
                            run6 = para2.add_run("\n" + Utils().getMyWord("{r_api_res}"))
                            run6.font.name = "Arial"
                            run6.font.size = Pt(10)
                            run6.font.bold = True
                            Creat_vuln_detail(self.projectTag).creat_table(document, info1, para3)

                # 处理文件上传漏洞
                elif vuln_info[3] == "upLoad":
                    sql = "select * from api_tree where id='%s'" % (vuln_info[1])
                    cursor.execute(sql)
                    api_infos = cursor.fetchall()
                    for api_info in api_infos:
                        para2 = Creat_vuln_detail(self.projectTag).insert_paragraph_after(para1)
                        para3 = para2.insert_paragraph_before("")
                        UserLogin = api_info[2]
                        api_path = api_info[1]
                        run = para3.add_run(
                            "2." + str(k - self.creat_num1 + 1) + " " + str(UserLogin) + Utils().getMyWord("{r_vuln_upload}") + "\n")
                        run.font.name = "Arial"
                        run.font.size = Pt(16)
                        run.font.bold = True
                        sql = "select path from js_file where id='%s'" % (vuln_info[2])
                        cursor.execute(sql)
                        js_paths = cursor.fetchall()
                        for js_path in js_paths:
                            run2 = para3.add_run(Utils().getMyWord("{r_api_addr}"))
                            run2.font.name = "Arial"
                            run2.font.size = Pt(10)
                            run2.font.bold = True
                            run3 = para3.add_run(api_path)
                            run3.font.name = "Arial"
                            run3.font.size = Pt(10)
                            run4 = para3.add_run("\n" + Utils().getMyWord("{r_api_js}"))
                            run4.font.name = "Arial"
                            run4.font.size = Pt(10)
                            run4.font.bold = True
                            run5 = para3.add_run(js_path[0])
                            run5.font.name = "Arial"
                            run5.font.size = Pt(10)
                            run5 = para3.add_run("\n" + Utils().getMyWord("{request_info}"))
                            run5.font.name = "Arial"
                            run5.font.size = Pt(10)
                            run5.font.bold = True
                            self.creat_num1 = self.creat_num1 + 1
                            Creat_vuln_detail(self.projectTag).creat_table(document, vuln_info[6], para2)
                            run6 = para2.add_run("\n" + Utils().getMyWord("{r_api_res}"))
                            run6.font.name = "Arial"
                            run6.font.size = Pt(10)
                            run6.font.bold = True
                            Creat_vuln_detail(self.projectTag).creat_table(document, vuln_info[5], para3)

                # 处理SQL注入漏洞
                elif vuln_info[3] == "SQL":
                    para2 = Creat_vuln_detail(self.projectTag).insert_paragraph_after(para1)
                    para3 = para2.insert_paragraph_before("")
                    UserLogin = api_info[2]
                    api_path = api_info[1]
                    run = para3.add_run(
                        "2." + str(k - self.creat_num1 + 1) + " " + str(UserLogin) + Utils().getMyWord("{r_vuln_sql}") + "\n")
                    run.font.name = "Arial"
                    run.font.size = Pt(16)
                    run.font.bold = True
                    sql = "select path from js_file where id='%s'" % (vuln_info[2])
                    cursor.execute(sql)
                    js_paths = cursor.fetchall()
                    for js_path in js_paths:
                        run2 = para3.add_run(Utils().getMyWord("{r_api_addr}"))
                        run2.font.name = "Arial"
                        run2.font.size = Pt(10)
                        run2.font.bold = True
                        run3 = para3.add_run(api_path)
                        run3.font.name = "Arial"
                        run3.font.size = Pt(10)
                        run4 = para3.add_run("\n" + Utils().getMyWord("{r_api_js}"))
                        run4.font.name = "Arial"
                        run4.font.size = Pt(10)
                        run4.font.bold = True
                        run5 = para3.add_run(js_path[0])
                        run5.font.name = "Arial"
                        run5.font.size = Pt(10)
                        run5 = para3.add_run("\n" + Utils().getMyWord("{request_info}"))
                        run5.font.name = "Arial"
                        run5.font.size = Pt(10)
                        run5.font.bold = True
                        self.creat_num1 = self.creat_num1 + 1
                        Creat_vuln_detail(self.projectTag).creat_table(document, vuln_info[6], para2)
                        run6 = para2.add_run("\n" + Utils().getMyWord("{r_api_res}"))
                        run6.font.name = "Arial"
                        run6.font.size = Pt(10)
                        run6.font.bold = True
                        Creat_vuln_detail(self.projectTag).creat_table(document, vuln_info[5], para3)

            self.log.debug("vuln_detail模块正常")
        except Exception as e:
            self.log.error("[Err] %s" % e)
