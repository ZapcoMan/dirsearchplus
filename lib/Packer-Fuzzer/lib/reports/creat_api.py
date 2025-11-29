#!/usr/bin/env python3
# -*- encoding: utf-8 -*-

import os,sqlite3,time,json
from docx import Document   #用来建立一个word对象
from docx.shared import Pt  #用来设置字体的大小
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from urllib.parse import urlparse
from lib.common.utils import Utils
from lib.Database import DatabaseType
from lib.common.CreatLog import creatLog
from lib.common.cmdline import CommandLines


class Creat_api():
    """
    用于创建API文档的类

    Attributes:
        projectTag (str): 项目标识符
        creat_num (int): 创建数量计数器
        log: 日志记录器实例
    """

    def __init__(self, projectTag):
        """
        初始化Creat_api类

        Args:
            projectTag (str): 项目的唯一标识符
        """
        self.projectTag = projectTag
        self.creat_num = 1
        self.log = creatLog().get_logger()

    def move_table_after(self, table, paragraph):
        """
        将表格移动到指定段落之后

        Args:
            table: 要移动的表格对象
            paragraph: 目标段落对象
        """
        tbl, p = table._tbl, paragraph._p
        p.addnext(tbl)

    def tabBgColor(self,table, cols, colorStr):
        """
        设置表格第一行的背景颜色

        Args:
            table: 表格对象
            cols (int): 列数
            colorStr (str): 颜色字符串（十六进制格式）
        """
        shading_list = locals()
        for i in range(cols):
            shading_list['shading_elm_' + str(i)] = parse_xml(
                r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStr))
            table.rows[0].cells[i]._tc.get_or_add_tcPr().append(shading_list['shading_elm_' + str(i)])

    def locat_api(self,document):
        """
        在文档中定位并处理{api_list}占位符

        Args:
            document: Word文档对象

        Returns:
            para1: 处理后的段落对象
        """
        for para in document.paragraphs:
            for i in range(len(para.runs)):
                if "{api_list}" in para.runs[i].text:
                    para.runs[i].text = para.runs[i].text.replace('{api_list}', '')
                    para1 = para.insert_paragraph_before("")
        return para1

    def creat_table(self,document,vuln_info,para2):
        """
        创建包含漏洞信息的表格

        Args:
            document: Word文档对象
            vuln_info (str): 漏洞信息内容
            para2: 段落对象
        """
        colorStr = 'F5F5F5'
        row = col = 1
        table = document.add_table(row, col)
        table.rows[0].cells[0].text = "%s" % (vuln_info)
        para3 = para2.insert_paragraph_before("")
        Creat_api(self.projectTag).tabBgColor(table, col, colorStr)
        Creat_api(self.projectTag).move_table_after(table, para2)

    def creat_api(self,document):
        """
        创建API列表文档内容

        Args:
            document: Word文档对象

        功能说明:
        - 定位API列表位置
        - 从数据库获取API树信息
        - 根据API信息生成相应的文档内容和表格
        - 处理API响应数据的JSON格式化显示
        """
        # 定位API列表在文档中的位置
        para1 = Creat_api(self.projectTag).locat_api(document)

        # 连接项目数据库获取API信息
        projectDBPath = DatabaseType(self.projectTag).getPathfromDB() + self.projectTag + ".db"
        connect = sqlite3.connect(os.sep.join(projectDBPath.split('/')))
        cursor = connect.cursor()
        connect.isolation_level = None
        sql = "select * from api_tree"
        cursor.execute(sql)
        api_infos = cursor.fetchall()

        # 遍历API信息并生成文档内容
        for api_info in api_infos:
            # 只处理状态为1或2的API
            if api_info[5] == 1 or api_info[5] == 2:
                para2 = para1.insert_paragraph_before("")
                api_path = api_info[1]

                # 获取关联的JS文件路径
                sql = "select path from js_file where id='%s'" % (api_info[6])
                cursor.execute(sql)
                js_paths = cursor.fetchall()

                for js_path in js_paths:
                    # 添加API地址信息
                    run2 = para2.add_run(Utils().getMyWord("{r_api_addr}") + "\n")
                    run2.font.name = "Arial"
                    run2.font.size = Pt(11)
                    run2.font.bold = True
                    run3 = para2.add_run(api_path)
                    run3.font.name = "Arial"
                    run3.font.size = Pt(11)

                    # 添加API相关JS文件信息
                    run4 = para2.add_run("\n" + Utils().getMyWord("{r_api_r_js}") + "\n")
                    run4.font.name = "Arial"
                    run4.font.size = Pt(11)
                    run4.font.bold = True
                    run5 = para2.add_run(js_path[0])
                    run5.font.name = "Arial"
                    run5.font.size = Pt(11)

                    # 添加API响应结果标题
                    run6 = para2.add_run("\n" + Utils().getMyWord("{r_api_res}") + "")
                    run6.font.name = "Arial"
                    run6.font.size = Pt(11)
                    run6.font.bold = True

                    # 尝试格式化并添加API响应数据
                    try:
                        if api_info[4] == None:
                            api_info1 = "\" \""
                            api_info_unicode = json.dumps(json.loads(api_info1),sort_keys=True, indent=4,ensure_ascii=False)
                            Creat_api(self.projectTag).creat_table(document, api_info_unicode, para2)
                        else:
                            api_info_unicode = json.dumps(json.loads(api_info[4]), sort_keys=True, indent = 4,ensure_ascii=False)
                            Creat_api(self.projectTag).creat_table(document, api_info_unicode, para2)
                        self.log.debug("api_info正常")
                    except Exception as e:
                        # 处理JSON解析失败的情况
                        if api_info[4] == None:
                            api_info1 = "\" \""
                            Creat_api(self.projectTag).creat_table(document, api_info1, para2)
                        else:
                            Creat_api(self.projectTag).creat_table(document, api_info[4], para2)
                        self.log.error("[Err] %s" % e)
